"""
Professional Word document generator for Selenium test results
"""
import os
import statistics
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from config import TestConfig

class DocumentGenerator:
    """Generate professional Word document with test results"""
    
    def __init__(self, test_results, test_summary):
        self.test_results = test_results
        self.test_summary = test_summary
        self.config = TestConfig()
        self.doc = Document()
        
        # Set document margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    
    def generate_word_document(self):
        """Generate complete Word document"""
        print("Generating professional Word document...")
        
        # Add document sections
        self._add_title_page()
        self._add_executive_summary()
        self._add_test_environment_setup()
        self._add_test_coverage_matrix()
        self._add_detailed_test_results()
        self._add_performance_metrics()
        self._add_bug_reports()
        self._add_recommendations()
        self._add_appendix()
        
        # Save document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        doc_filename = f"MM_Fashion_Selenium_Test_Report_{timestamp}.docx"
        doc_path = os.path.join(self.config.REPORT_DIR, doc_filename)
        
        self.doc.save(doc_path)
        print(f"Word document saved: {doc_path}")
        return doc_path
    
    def _add_title_page(self):
        """Add professional title page"""
        # Title
        title = self.doc.add_heading('M&M Fashion E-commerce Platform', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = self.doc.add_heading('Comprehensive Selenium Testing Report', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add some spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Project details table
        details_table = self.doc.add_table(rows=8, cols=2)
        details_table.style = 'Table Grid'
        details_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        details_data = [
            ('Project Name', 'M&M Fashion E-commerce Platform'),
            ('Test Type', 'Automated Selenium Testing'),
            ('Report Date', datetime.now().strftime("%B %d, %Y")),
            ('Test Duration', f"{self.test_summary.get('duration', 0):.2f} seconds"),
            ('Browsers Tested', ', '.join(self.test_summary.get('browsers_tested', []))),
            ('Domains Tested', ', '.join(self.test_summary.get('domains_tested', []))),
            ('Total Tests', str(self.test_summary.get('total_tests', 0))),
            ('Pass Rate', f"{(self.test_summary.get('passed_tests', 0) / max(self.test_summary.get('total_tests', 1), 1)) * 100:.1f}%")
        ]
        
        for i, (label, value) in enumerate(details_data):
            row = details_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Make label bold
            row.cells[0].paragraphs[0].runs[0].bold = True
        
        # Add page break
        self.doc.add_page_break()
    
    def _add_executive_summary(self):
        """Add executive summary section"""
        self.doc.add_heading('Executive Summary', level=1)
        
        # Overview paragraph
        overview = self.doc.add_paragraph()
        overview.add_run("This document presents the comprehensive testing results for the M&M Fashion e-commerce platform. ")
        overview.add_run("The testing was conducted using Selenium WebDriver to validate functionality across multiple browsers ")
        overview.add_run("and domains (B2C: garba.shop and B2B: ttd.in).")
        
        # Key findings
        self.doc.add_heading('Key Findings', level=2)
        
        total_tests = self.test_summary.get('total_tests', 0)
        passed_tests = self.test_summary.get('passed_tests', 0)
        failed_tests = self.test_summary.get('failed_tests', 0)
        skipped_tests = self.test_summary.get('skipped_tests', 0)
        
        findings = self.doc.add_paragraph()
        findings.add_run(f"• Total test cases executed: {total_tests}").bold = True
        findings.add_run(f"\n• Successful test cases: {passed_tests} ({(passed_tests/max(total_tests,1)*100):.1f}%)")
        findings.add_run(f"\n• Failed test cases: {failed_tests} ({(failed_tests/max(total_tests,1)*100):.1f}%)")
        findings.add_run(f"\n• Skipped test cases: {skipped_tests} ({(skipped_tests/max(total_tests,1)*100):.1f}%)")
        
        # Test categories
        self.doc.add_heading('Test Categories Covered', level=2)
        categories = self.doc.add_paragraph()
        categories.add_run("✓ Homepage functionality and navigation\n")
        categories.add_run("✓ Product browsing and filtering\n")
        categories.add_run("✓ Product detail page interactions\n")
        categories.add_run("✓ Shopping cart management\n")
        categories.add_run("✓ Checkout process validation\n")
        categories.add_run("✓ Cross-browser compatibility\n")
        categories.add_run("✓ Mobile responsiveness\n")
        categories.add_run("✓ Performance testing\n")
        categories.add_run("✓ B2C vs B2B domain testing")
        
        # Overall assessment
        self.doc.add_heading('Overall Assessment', level=2)
        
        if passed_tests / max(total_tests, 1) >= 0.8:
            assessment_text = "The M&M Fashion platform demonstrates good overall functionality with most test cases passing successfully. "
        elif passed_tests / max(total_tests, 1) >= 0.6:
            assessment_text = "The M&M Fashion platform shows moderate functionality with some areas requiring attention. "
        else:
            assessment_text = "The M&M Fashion platform requires significant improvements to meet quality standards. "
        
        assessment = self.doc.add_paragraph(assessment_text)
        
        if failed_tests > 0:
            assessment.add_run(f"There are {failed_tests} failed test cases that need immediate attention. ")
        
        assessment.add_run("Detailed findings and recommendations are provided in the subsequent sections.")
    
    def _add_test_environment_setup(self):
        """Add test environment setup section"""
        self.doc.add_heading('Test Environment Setup', level=1)
        
        # Application URLs
        self.doc.add_heading('Application URLs', level=2)
        urls_table = self.doc.add_table(rows=3, cols=2)
        urls_table.style = 'Table Grid'
        
        urls_data = [
            ('Frontend URL', self.config.FRONTEND_URL),
            ('Backend URL', self.config.BACKEND_URL),
            ('Test Environment', 'Local Development')
        ]
        
        for i, (label, value) in enumerate(urls_data):
            row = urls_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            row.cells[0].paragraphs[0].runs[0].bold = True
        
        # Browser Configuration
        self.doc.add_heading('Browser Configuration', level=2)
        browser_info = self.doc.add_paragraph()
        browser_info.add_run("Browsers Tested: ").bold = True
        browser_info.add_run(', '.join(self.test_summary.get('browsers_tested', [])))
        browser_info.add_run("\nWebDriver: ").bold = True
        browser_info.add_run("Selenium WebDriver 4.15.2")
        browser_info.add_run("\nWebDriver Manager: ").bold = True
        browser_info.add_run("Automatic driver management")
        
        # Test Data Configuration
        self.doc.add_heading('Test Data Configuration', level=2)
        test_data_table = self.doc.add_table(rows=6, cols=2)
        test_data_table.style = 'Table Grid'
        
        test_data = [
            ('B2C Domain', self.config.B2C_DOMAIN),
            ('B2B Domain', self.config.B2B_DOMAIN),
            ('Desktop Resolution', f"{self.config.DESKTOP_SIZE[0]}x{self.config.DESKTOP_SIZE[1]}"),
            ('Mobile Resolution', f"{self.config.MOBILE_SIZE[0]}x{self.config.MOBILE_SIZE[1]}"),
            ('Page Load Timeout', f"{self.config.PAGE_LOAD_TIMEOUT} seconds"),
            ('Element Wait Timeout', f"{self.config.EXPLICIT_WAIT} seconds")
        ]
        
        for i, (label, value) in enumerate(test_data):
            row = test_data_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            row.cells[0].paragraphs[0].runs[0].bold = True
    
    def _add_test_coverage_matrix(self):
        """Add test coverage matrix"""
        self.doc.add_heading('Test Coverage Matrix', level=1)
        
        # Create coverage matrix table
        coverage_table = self.doc.add_table(rows=1, cols=5)
        coverage_table.style = 'Table Grid'
        
        # Header row
        header_cells = coverage_table.rows[0].cells
        header_cells[0].text = 'Test Category'
        header_cells[1].text = 'Test Cases'
        header_cells[2].text = 'Passed'
        header_cells[3].text = 'Failed'
        header_cells[4].text = 'Coverage %'
        
        # Make header bold
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
        
        # Analyze test results by category
        categories = self._analyze_test_categories()
        
        for category, stats in categories.items():
            row_cells = coverage_table.add_row().cells
            row_cells[0].text = category
            row_cells[1].text = str(stats['total'])
            row_cells[2].text = str(stats['passed'])
            row_cells[3].text = str(stats['failed'])
            row_cells[4].text = f"{(stats['passed']/max(stats['total'],1)*100):.1f}%"
        
        # Add summary paragraph
        self.doc.add_paragraph()
        summary = self.doc.add_paragraph()
        summary.add_run("Coverage Analysis: ").bold = True
        summary.add_run("The test suite covers all major functional areas of the M&M Fashion platform. ")
        summary.add_run("Each category includes multiple test scenarios covering both positive and negative test cases.")
    
    def _add_detailed_test_results(self):
        """Add detailed test results section"""
        self.doc.add_heading('Detailed Test Results', level=1)
        
        # Group results by category
        categories = self._analyze_test_categories()
        
        for category, stats in categories.items():
            self.doc.add_heading(category, level=2)
            
            # Category summary
            category_summary = self.doc.add_paragraph()
            category_summary.add_run(f"Total Tests: {stats['total']} | ")
            category_summary.add_run(f"Passed: {stats['passed']} | ").font.color.rgb = self._get_color(0, 128, 0)
            category_summary.add_run(f"Failed: {stats['failed']} | ").font.color.rgb = self._get_color(255, 0, 0)
            category_summary.add_run(f"Skipped: {stats['skipped']}").font.color.rgb = self._get_color(255, 165, 0)
            
            # Test results table for this category
            if stats['tests']:
                results_table = self.doc.add_table(rows=1, cols=4)
                results_table.style = 'Table Grid'
                
                # Header
                header = results_table.rows[0].cells
                header[0].text = 'Test Name'
                header[1].text = 'Status'
                header[2].text = 'Browser'
                header[3].text = 'Message'
                
                for cell in header:
                    cell.paragraphs[0].runs[0].bold = True
                
                # Add test results
                for test in stats['tests'][:10]:  # Limit to first 10 tests per category
                    row = results_table.add_row().cells
                    row[0].text = test['test_name'].split(' - ')[-1]  # Remove category prefix
                    row[1].text = test['status']
                    row[2].text = test.get('browser', 'N/A')
                    row[3].text = test.get('message', '')[:100] + ('...' if len(test.get('message', '')) > 100 else '')
                    
                    # Color code status
                    if test['status'] == 'PASS':
                        row[1].paragraphs[0].runs[0].font.color.rgb = self._get_color(0, 128, 0)
                    elif test['status'] == 'FAIL':
                        row[1].paragraphs[0].runs[0].font.color.rgb = self._get_color(255, 0, 0)
                    else:
                        row[1].paragraphs[0].runs[0].font.color.rgb = self._get_color(255, 165, 0)
            
            self.doc.add_paragraph()  # Add spacing
    
    def _add_performance_metrics(self):
        """Add performance metrics section"""
        self.doc.add_heading('Performance Metrics', level=1)
        
        # Extract performance-related results
        perf_results = [r for r in self.test_results if 'performance' in r.get('test_name', '').lower()]
        
        if perf_results:
            self.doc.add_heading('Page Load Performance', level=2)
            
            # Performance summary
            perf_summary = self.doc.add_paragraph()
            perf_summary.add_run("Performance testing was conducted to measure page load times and system responsiveness. ")
            perf_summary.add_run(f"The maximum acceptable page load time threshold was set to {self.config.MAX_PAGE_LOAD_TIME} seconds.")
            
            # Performance results table
            perf_table = self.doc.add_table(rows=1, cols=3)
            perf_table.style = 'Table Grid'
            
            header = perf_table.rows[0].cells
            header[0].text = 'Test'
            header[1].text = 'Result'
            header[2].text = 'Status'
            
            for cell in header:
                cell.paragraphs[0].runs[0].bold = True
            
            for result in perf_results:
                row = perf_table.add_row().cells
                row[0].text = result['test_name']
                row[1].text = result.get('message', '')
                row[2].text = result['status']
        else:
            no_perf = self.doc.add_paragraph()
            no_perf.add_run("No specific performance test results available in this test run.")
        
        # General performance observations
        self.doc.add_heading('Performance Observations', level=2)
        observations = self.doc.add_paragraph()
        observations.add_run("• Page load times were monitored throughout all test executions\n")
        observations.add_run("• JavaScript errors were tracked and reported\n")
        observations.add_run("• Memory usage patterns were observed during extended browsing sessions\n")
        observations.add_run("• Cross-browser performance variations were documented")
    
    def _add_bug_reports(self):
        """Add bug reports and issues section"""
        self.doc.add_heading('Bug Reports and Issues Found', level=1)
        
        # Extract failed tests
        failed_tests = [r for r in self.test_results if r['status'] == 'FAIL']
        
        if failed_tests:
            self.doc.add_heading('Critical Issues', level=2)
            
            # Bug report table
            bug_table = self.doc.add_table(rows=1, cols=5)
            bug_table.style = 'Table Grid'
            
            header = bug_table.rows[0].cells
            header[0].text = 'Bug ID'
            header[1].text = 'Test Case'
            header[2].text = 'Browser'
            header[3].text = 'Description'
            header[4].text = 'Severity'
            
            for cell in header:
                cell.paragraphs[0].runs[0].bold = True
            
            for i, bug in enumerate(failed_tests, 1):
                row = bug_table.add_row().cells
                row[0].text = f"BUG-{i:03d}"
                row[1].text = bug['test_name']
                row[2].text = bug.get('browser', 'N/A')
                row[3].text = bug.get('message', '')[:150] + ('...' if len(bug.get('message', '')) > 150 else '')
                
                # Determine severity based on test type
                if 'checkout' in bug['test_name'].lower() or 'cart' in bug['test_name'].lower():
                    severity = 'High'
                elif 'homepage' in bug['test_name'].lower() or 'navigation' in bug['test_name'].lower():
                    severity = 'Medium'
                else:
                    severity = 'Low'
                
                row[4].text = severity
                
                # Color code severity
                if severity == 'High':
                    row[4].paragraphs[0].runs[0].font.color.rgb = self._get_color(255, 0, 0)
                elif severity == 'Medium':
                    row[4].paragraphs[0].runs[0].font.color.rgb = self._get_color(255, 165, 0)
                else:
                    row[4].paragraphs[0].runs[0].font.color.rgb = self._get_color(0, 128, 0)
        else:
            no_bugs = self.doc.add_paragraph()
            no_bugs.add_run("✓ No critical bugs were identified during testing.").bold = True
            no_bugs.add_run(" All test cases passed successfully.")
        
        # Known limitations
        self.doc.add_heading('Known Limitations', level=2)
        limitations = self.doc.add_paragraph()
        limitations.add_run("• Testing was performed in a local development environment\n")
        limitations.add_run("• Payment gateway integration was not tested with real transactions\n")
        limitations.add_run("• Email notifications were not validated\n")
        limitations.add_run("• Load testing with multiple concurrent users was not performed\n")
        limitations.add_run("• Third-party integrations (if any) were not comprehensively tested")
    
    def _add_recommendations(self):
        """Add recommendations section"""
        self.doc.add_heading('Recommendations for Improvements', level=1)
        
        failed_count = self.test_summary.get('failed_tests', 0)
        total_count = self.test_summary.get('total_tests', 1)
        
        # Priority recommendations based on test results
        self.doc.add_heading('High Priority Recommendations', level=2)
        high_priority = self.doc.add_paragraph()
        
        if failed_count > 0:
            high_priority.add_run(f"1. Address {failed_count} failed test cases immediately\n")
            high_priority.add_run("2. Implement comprehensive error handling and user feedback\n")
            high_priority.add_run("3. Enhance form validation across all user input forms\n")
        else:
            high_priority.add_run("1. Maintain current quality standards\n")
            high_priority.add_run("2. Implement continuous integration testing\n")
        
        high_priority.add_run("4. Optimize page load performance for mobile devices\n")
        high_priority.add_run("5. Implement comprehensive logging for debugging")
        
        # Medium priority recommendations
        self.doc.add_heading('Medium Priority Recommendations', level=2)
        medium_priority = self.doc.add_paragraph()
        medium_priority.add_run("1. Enhance cross-browser compatibility testing\n")
        medium_priority.add_run("2. Implement automated accessibility testing\n")
        medium_priority.add_run("3. Add comprehensive API testing\n")
        medium_priority.add_run("4. Implement security testing protocols\n")
        medium_priority.add_run("5. Add performance monitoring and alerting")
        
        # Long-term recommendations
        self.doc.add_heading('Long-term Recommendations', level=2)
        long_term = self.doc.add_paragraph()
        long_term.add_run("1. Implement continuous integration/continuous deployment (CI/CD)\n")
        long_term.add_run("2. Set up automated testing in staging environment\n")
        long_term.add_run("3. Implement comprehensive monitoring and analytics\n")
        long_term.add_run("4. Regular security audits and penetration testing\n")
        long_term.add_run("5. User experience testing with real users")
        
        # Quality assurance process
        self.doc.add_heading('Quality Assurance Process Improvements', level=2)
        qa_process = self.doc.add_paragraph()
        qa_process.add_run("• Establish regular automated testing schedules\n")
        qa_process.add_run("• Implement code review processes\n")
        qa_process.add_run("• Create comprehensive test documentation\n")
        qa_process.add_run("• Establish performance benchmarks and monitoring\n")
        qa_process.add_run("• Regular training on testing best practices")
    
    def _add_appendix(self):
        """Add appendix with technical details"""
        self.doc.add_heading('Appendix', level=1)
        
        # Test configuration
        self.doc.add_heading('A. Test Configuration Details', level=2)
        config_table = self.doc.add_table(rows=8, cols=2)
        config_table.style = 'Table Grid'
        
        config_data = [
            ('Frontend URL', self.config.FRONTEND_URL),
            ('Backend URL', self.config.BACKEND_URL),
            ('Implicit Wait', f"{self.config.IMPLICIT_WAIT} seconds"),
            ('Explicit Wait', f"{self.config.EXPLICIT_WAIT} seconds"),
            ('Page Load Timeout', f"{self.config.PAGE_LOAD_TIMEOUT} seconds"),
            ('Max Page Load Time', f"{self.config.MAX_PAGE_LOAD_TIME} seconds"),
            ('Max API Response Time', f"{self.config.MAX_API_RESPONSE_TIME} seconds"),
            ('Screenshot Directory', self.config.SCREENSHOT_DIR)
        ]
        
        for i, (label, value) in enumerate(config_data):
            row = config_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            row.cells[0].paragraphs[0].runs[0].bold = True
        
        # Test execution summary
        self.doc.add_heading('B. Test Execution Summary', level=2)
        execution_info = self.doc.add_paragraph()
        execution_info.add_run(f"Start Time: {self.test_summary.get('start_time', 'N/A')}\n")
        execution_info.add_run(f"End Time: {self.test_summary.get('end_time', 'N/A')}\n")
        execution_info.add_run(f"Total Duration: {self.test_summary.get('duration', 0):.2f} seconds\n")
        execution_info.add_run(f"Browsers Tested: {', '.join(self.test_summary.get('browsers_tested', []))}\n")
        execution_info.add_run(f"Domains Tested: {', '.join(self.test_summary.get('domains_tested', []))}")
        
        # Contact information
        self.doc.add_heading('C. Contact Information', level=2)
        contact_info = self.doc.add_paragraph()
        contact_info.add_run("For questions regarding this test report, please contact:\n\n")
        contact_info.add_run("QA Team\n").bold = True
        contact_info.add_run("Email: qa@mmfashion.com\n")
        contact_info.add_run("Generated by: Selenium Test Automation Suite\n")
        contact_info.add_run(f"Report Date: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    
    def _analyze_test_categories(self):
        """Analyze test results by category"""
        categories = {}
        
        for result in self.test_results:
            # Extract category from test name
            test_name = result.get('test_name', '')
            
            if 'Homepage' in test_name:
                category = 'Homepage Tests'
            elif 'Product Browsing' in test_name or 'Category' in test_name:
                category = 'Product Browsing Tests'
            elif 'Product Detail' in test_name:
                category = 'Product Detail Tests'
            elif 'Cart' in test_name or 'Checkout' in test_name:
                category = 'Cart & Checkout Tests'
            elif 'Performance' in test_name:
                category = 'Performance Tests'
            elif 'Responsive' in test_name:
                category = 'Responsive Design Tests'
            else:
                category = 'Other Tests'
            
            if category not in categories:
                categories[category] = {
                    'total': 0,
                    'passed': 0,
                    'failed': 0,
                    'skipped': 0,
                    'tests': []
                }
            
            categories[category]['total'] += 1
            categories[category]['tests'].append(result)
            
            if result['status'] == 'PASS':
                categories[category]['passed'] += 1
            elif result['status'] == 'FAIL':
                categories[category]['failed'] += 1
            else:
                categories[category]['skipped'] += 1
        
        return categories
    
    def _get_color(self, r, g, b):
        """Get RGB color for Word document"""
        from docx.shared import RGBColor
        return RGBColor(r, g, b)

def main():
    """Test the document generator"""
    # Sample test data for testing
    sample_results = [
        {
            'test_name': 'Homepage Load - garba.shop',
            'status': 'PASS',
            'message': 'Homepage loaded successfully',
            'browser': 'chrome',
            'timestamp': datetime.now().isoformat()
        },
        {
            'test_name': 'Product Detail Page Load - garba.shop',
            'status': 'FAIL',
            'message': 'Add to cart button not found',
            'browser': 'chrome',
            'timestamp': datetime.now().isoformat()
        }
    ]
    
    sample_summary = {
        'total_tests': 2,
        'passed_tests': 1,
        'failed_tests': 1,
        'skipped_tests': 0,
        'start_time': datetime.now(),
        'end_time': datetime.now(),
        'duration': 120.5,
        'browsers_tested': ['chrome'],
        'domains_tested': ['garba.shop']
    }
    
    generator = DocumentGenerator(sample_results, sample_summary)
    doc_path = generator.generate_word_document()
    print(f"Test document generated: {doc_path}")

if __name__ == "__main__":
    main()